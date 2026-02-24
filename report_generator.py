"""
AI_AUDITOR - Report Generator
Génère rapports professionnels en PDF, Word, Excel

Edge-first: Tout généré LOCALEMENT, aucun appel API
"""

from typing import List, Dict, Optional
from datetime import datetime
import io
import pandas as pd
from pathlib import Path

# PDF generation
from reportlab.lib.pagesizes import letter, A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import (
    SimpleDocTemplate, Table, TableStyle, Paragraph, 
    PageBreak, Spacer, Image, PageTemplate, Frame
)
from reportlab.lib import colors
from reportlab.pdfgen import canvas
from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT, TA_JUSTIFY

# Word generation
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Excel generation
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Border, Side, Alignment
from openpyxl.utils import get_column_letter


class ProfessionalAuditReport:
    """Générateur de rapports d'audit professionnels"""
    
    def __init__(self, company_name: str = "AI_Auditor", company_logo: Optional[str] = None):
        """
        Initialiser le générateur
        
        Args:
            company_name: Nom du cabinet/entreprise
            company_logo: Chemin au logo (optionnel)
        """
        self.company_name = company_name
        self.company_logo = company_logo
        self.timestamp = datetime.now()
    
    def generate_from_anomalies(
        self,
        client_name: str,
        anomalies: List[Dict],
        format: str = "PDF"
    ) -> bytes:
        """
        Générer rapport à partir des anomalies
        
        Args:
            client_name: Nom du client audité
            anomalies: Liste des anomalies détectées
            format: "PDF", "Word" ou "Excel"
        
        Returns:
            Bytes du fichier rapport
        """
        if format.upper() == "PDF":
            return self._generate_pdf(client_name, anomalies)
        elif format.upper() == "WORD":
            return self._generate_word(client_name, anomalies)
        elif format.upper() == "EXCEL":
            return self._generate_excel(client_name, anomalies)
        else:
            raise ValueError(f"Format non supporté: {format}")
    
    # ═══════════════════════════════════════════════════════════════
    #                       PDF GENERATION
    # ═══════════════════════════════════════════════════════════════
    
    def _generate_pdf(self, client_name: str, anomalies: List[Dict]) -> bytes:
        """Générer rapport PDF professionnel"""
        
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(
            buffer,
            pagesize=letter,
            rightMargin=0.75 * inch,
            leftMargin=0.75 * inch,
            topMargin=1 * inch,
            bottomMargin=0.75 * inch,
            title=f"Rapport d'Audit - {client_name}"
        )
        
        # Build story
        story = []
        styles = getSampleStyleSheet()
        
        # Styles personnalisés
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=24,
            textColor=colors.HexColor('#1a1a2e'),
            spaceAfter=12,
            alignment=TA_CENTER,
            fontName='Helvetica-Bold'
        )
        
        heading_style = ParagraphStyle(
            'CustomHeading',
            parent=styles['Heading2'],
            fontSize=14,
            textColor=colors.HexColor('#0066cc'),
            spaceAfter=6,
            spaceBefore=12,
            fontName='Helvetica-Bold'
        )
        
        normal_style = ParagraphStyle(
            'CustomNormal',
            parent=styles['Normal'],
            fontSize=11,
            textColor=colors.HexColor('#333333'),
            spaceAfter=6,
            fontName='Helvetica'
        )
        
        # En-tête
        story.append(Paragraph(f"<b>{self.company_name}</b>", heading_style))
        story.append(Paragraph("Rapport d'Audit Fiscal", title_style))
        story.append(Spacer(1, 0.3 * inch))
        
        # Infos rapport
        info_data = [
            ['Client:', client_name],
            ['Date d\'audit:', self.timestamp.strftime('%d-%m-%Y %H:%M')],
            ['Audité par:', self.company_name],
            ['Confiance moyenne:', f"{self._avg_confidence(anomalies):.1f}%"]
        ]
        
        info_table = Table(info_data, colWidths=[1.5 * inch, 3.5 * inch])
        info_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (0, -1), colors.HexColor('#f0f0f0')),
            ('TEXTCOLOR', (0, 0), (-1, -1), colors.black),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, -1), 10),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
            ('GRID', (0, 0), (-1, -1), 1, colors.grey)
        ]))
        
        story.append(info_table)
        story.append(Spacer(1, 0.3 * inch))
        
        # Résumé exécutif
        story.append(Paragraph("<b>RÉSUMÉ EXÉCUTIF</b>", heading_style))
        
        summary_data = [
            ['Métrique', 'Valeur'],
            ['Anomalies détectées', str(len(anomalies))],
            ['Impact total estimé', f"${sum(a.get('Impact_Estimation', 0) for a in anomalies):,.2f}"],
            ['Alertes critiques', str(sum(1 for a in anomalies if a.get('Risque') == 'CRITIQUE'))],
            ['Alertes moyennes', str(sum(1 for a in anomalies if a.get('Risque') == 'MOYEN'))]
        ]
        
        summary_table = Table(summary_data, colWidths=[2.5 * inch, 3.5 * inch])
        summary_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#0066cc')),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, -1), 10),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
            ('GRID', (0, 0), (-1, -1), 1, colors.black),
            ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#f9f9f9')])
        ]))
        
        story.append(summary_table)
        story.append(Spacer(1, 0.2 * inch))
        
        # Distribution par type
        story.append(Paragraph("<b>Distribution des anomalies par type</b>", heading_style))
        
        type_counts = {}
        for anom in anomalies:
            anom_type = anom.get('Type', 'Inconnu')
            type_counts[anom_type] = type_counts.get(anom_type, 0) + 1
        
        type_data = [['Type', 'Nombre', 'Pourcentage']]
        for anom_type, count in sorted(type_counts.items(), key=lambda x: x[1], reverse=True):
            percentage = (count / len(anomalies)) * 100
            type_data.append([anom_type, str(count), f"{percentage:.1f}%"])
        
        type_table = Table(type_data, colWidths=[2.5 * inch, 1.5 * inch, 1.5 * inch])
        type_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#0066cc')),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, -1), 9),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
            ('GRID', (0, 0), (-1, -1), 1, colors.black),
            ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#f9f9f9')])
        ]))
        
        story.append(type_table)
        story.append(PageBreak())
        
        # Détail anomalies
        story.append(Paragraph("<b>DÉTAIL DES ANOMALIES</b>", heading_style))
        
        for idx, anomaly in enumerate(sorted(anomalies, key=lambda x: x.get('Impact_Estimation', 0), reverse=True)):
            story.append(Paragraph(f"<b>{idx+1}. {anomaly.get('Type', 'Anomalie')}</b>", heading_style))
            
            detail_data = [
                ['Description:', anomaly.get('Description', 'N/A')],
                ['Fournisseur:', anomaly.get('Fournisseur', 'N/A')],
                ['Montant:', f"${anomaly.get('Montant', 0):,.2f}"],
                ['Impact estimé:', f"${anomaly.get('Impact_Estimation', 0):,.2f}"],
                ['Niveau de risque:', anomaly.get('Risque', 'N/A')],
                ['Confiance:', f"{anomaly.get('Confiance', 0):.1f}%"],
                ['Recommandation:', anomaly.get('Recommandation', 'N/A')]
            ]
            
            detail_table = Table(detail_data, colWidths=[1.5 * inch, 4.5 * inch])
            detail_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (0, -1), colors.HexColor('#e8e8e8')),
                ('TEXTCOLOR', (0, 0), (-1, -1), colors.black),
                ('ALIGN', (0, 0), (0, -1), 'LEFT'),
                ('ALIGN', (1, 0), (1, -1), 'LEFT'),
                ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, -1), 9),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
                ('GRID', (0, 0), (-1, -1), 1, colors.grey),
                ('VALIGN', (0, 0), (-1, -1), 'TOP')
            ]))
            
            story.append(detail_table)
            story.append(Spacer(1, 0.15 * inch))
        
        # Footer
        story.append(Spacer(1, 0.3 * inch))
        story.append(Paragraph(
            f"<i>Ce rapport a été généré automatiquement par {self.company_name}. "
            f"Données 100% locales, conforme LPRPDE.</i>",
            normal_style
        ))
        
        # Build PDF
        doc.build(story)
        buffer.seek(0)
        return buffer.getvalue()
    
    # ═══════════════════════════════════════════════════════════════
    #                       WORD GENERATION
    # ═══════════════════════════════════════════════════════════════
    
    def _generate_word(self, client_name: str, anomalies: List[Dict]) -> bytes:
        """Générer rapport Word professionnel (éditable)"""
        
        doc = Document()
        
        # Style document
        style = doc.styles['Normal']
        style.font.name = 'Calibri'
        style.font.size = Pt(11)
        
        # En-tête
        title = doc.add_heading(f"{self.company_name}", level=0)
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        title_run = title.runs[0]
        title_run.font.size = Pt(18)
        title_run.font.color.rgb = RGBColor(26, 26, 46)
        
        report_title = doc.add_heading("RAPPORT D'AUDIT FISCAL", level=1)
        report_title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        # Infos rapport
        doc.add_heading("Informations du rapport", level=2)
        
        table = doc.add_table(rows=4, cols=2)
        table.style = 'Light Grid Accent 1'
        
        rows = table.rows
        rows[0].cells[0].text = "Client:"
        rows[0].cells[1].text = client_name
        rows[1].cells[0].text = "Date d'audit:"
        rows[1].cells[1].text = self.timestamp.strftime('%d-%m-%Y %H:%M')
        rows[2].cells[0].text = "Audité par:"
        rows[2].cells[1].text = self.company_name
        rows[3].cells[0].text = "Confiance moyenne:"
        rows[3].cells[1].text = f"{self._avg_confidence(anomalies):.1f}%"
        
        doc.add_paragraph()
        
        # Résumé exécutif
        doc.add_heading("Résumé exécutif", level=2)
        
        summary_para = doc.add_paragraph()
        summary_para.add_run(f"Nombre d'anomalies détectées: ").bold = True
        summary_para.add_run(f"{len(anomalies)}\n")
        
        summary_para.add_run(f"Impact total estimé: ").bold = True
        summary_para.add_run(f"${sum(a.get('Impact_Estimation', 0) for a in anomalies):,.2f}\n")
        
        summary_para.add_run(f"Alertes critiques: ").bold = True
        summary_para.add_run(f"{sum(1 for a in anomalies if a.get('Risque') == 'CRITIQUE')}\n")
        
        summary_para.add_run(f"Alertes moyennes: ").bold = True
        summary_para.add_run(f"{sum(1 for a in anomalies if a.get('Risque') == 'MOYEN')}")
        
        doc.add_paragraph()
        
        # Distribution
        doc.add_heading("Distribution des anomalies", level=2)
        
        type_counts = {}
        for anom in anomalies:
            anom_type = anom.get('Type', 'Inconnu')
            type_counts[anom_type] = type_counts.get(anom_type, 0) + 1
        
        for anom_type, count in sorted(type_counts.items(), key=lambda x: x[1], reverse=True):
            percentage = (count / len(anomalies)) * 100
            doc.add_paragraph(f"{anom_type}: {count} ({percentage:.1f}%)", style='List Bullet')
        
        doc.add_page_break()
        
        # Détail anomalies
        doc.add_heading("Détail des anomalies", level=2)
        
        for idx, anomaly in enumerate(sorted(anomalies, key=lambda x: x.get('Impact_Estimation', 0), reverse=True)):
            doc.add_heading(f"{idx+1}. {anomaly.get('Type', 'Anomalie')}", level=3)
            
            detail = doc.add_paragraph()
            detail.add_run("Description: ").bold = True
            detail.add_run(f"{anomaly.get('Description', 'N/A')}\n")
            
            detail.add_run("Fournisseur: ").bold = True
            detail.add_run(f"{anomaly.get('Fournisseur', 'N/A')}\n")
            
            detail.add_run("Montant: ").bold = True
            detail.add_run(f"${anomaly.get('Montant', 0):,.2f}\n")
            
            detail.add_run("Impact estimé: ").bold = True
            detail.add_run(f"${anomaly.get('Impact_Estimation', 0):,.2f}\n")
            
            detail.add_run("Niveau de risque: ").bold = True
            detail.add_run(f"{anomaly.get('Risque', 'N/A')}\n")
            
            detail.add_run("Confiance: ").bold = True
            detail.add_run(f"{anomaly.get('Confiance', 0):.1f}%\n")
            
            detail.add_run("Recommandation: ").bold = True
            detail.add_run(f"{anomaly.get('Recommandation', 'N/A')}")
            
            doc.add_paragraph()
        
        # Footer
        doc.add_paragraph()
        footer = doc.add_paragraph()
        footer.add_run(
            f"Ce rapport a été généré automatiquement par {self.company_name}. "
            f"Données 100% locales, conforme LPRPDE."
        ).italic = True
        
        # Save to bytes
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()
    
    # ═══════════════════════════════════════════════════════════════
    #                       EXCEL GENERATION
    # ═══════════════════════════════════════════════════════════════
    
    def _generate_excel(self, client_name: str, anomalies: List[Dict]) -> bytes:
        """Générer rapport Excel avec données détaillées"""
        
        wb = Workbook()
        
        # Sheet 1: Résumé
        ws_summary = wb.active
        ws_summary.title = "Résumé"
        
        # Headers
        header_fill = PatternFill(start_color="0066CC", end_color="0066CC", fill_type="solid")
        header_font = Font(bold=True, color="FFFFFF", size=12)
        
        ws_summary['A1'] = "RAPPORT D'AUDIT FISCAL"
        ws_summary['A1'].font = Font(bold=True, size=14)
        
        ws_summary['A3'] = "Client:"
        ws_summary['B3'] = client_name
        
        ws_summary['A4'] = "Date d'audit:"
        ws_summary['B4'] = self.timestamp.strftime('%d-%m-%Y %H:%M')
        
        ws_summary['A5'] = "Audité par:"
        ws_summary['B5'] = self.company_name
        
        ws_summary['A7'] = "Métrique"
        ws_summary['B7'] = "Valeur"
        
        for cell in ['A7', 'B7']:
            ws_summary[cell].fill = header_fill
            ws_summary[cell].font = header_font
        
        ws_summary['A8'] = "Anomalies détectées"
        ws_summary['B8'] = len(anomalies)
        
        ws_summary['A9'] = "Impact total estimé"
        ws_summary['B9'] = sum(a.get('Impact_Estimation', 0) for a in anomalies)
        
        ws_summary['A10'] = "Alertes critiques"
        ws_summary['B10'] = sum(1 for a in anomalies if a.get('Risque') == 'CRITIQUE')
        
        ws_summary['A11'] = "Confiance moyenne"
        ws_summary['B11'] = self._avg_confidence(anomalies)
        
        ws_summary.column_dimensions['A'].width = 25
        ws_summary.column_dimensions['B'].width = 25
        
        # Sheet 2: Détail anomalies
        ws_details = wb.create_sheet("Anomalies")
        
        headers = ['Type', 'Description', 'Fournisseur', 'Montant', 'Impact', 'Risque', 'Confiance', 'Recommandation']
        ws_details.append(headers)
        
        for cell in ws_details[1]:
            cell.fill = header_fill
            cell.font = header_font
            cell.alignment = Alignment(horizontal='center')
        
        for anomaly in sorted(anomalies, key=lambda x: x.get('Impact_Estimation', 0), reverse=True):
            ws_details.append([
                anomaly.get('Type', ''),
                anomaly.get('Description', ''),
                anomaly.get('Fournisseur', ''),
                anomaly.get('Montant', 0),
                anomaly.get('Impact_Estimation', 0),
                anomaly.get('Risque', ''),
                anomaly.get('Confiance', 0),
                anomaly.get('Recommandation', '')
            ])
        
        # Format columns
        ws_details.column_dimensions['A'].width = 20
        ws_details.column_dimensions['B'].width = 30
        ws_details.column_dimensions['C'].width = 20
        ws_details.column_dimensions['D'].width = 15
        ws_details.column_dimensions['E'].width = 15
        ws_details.column_dimensions['F'].width = 12
        ws_details.column_dimensions['G'].width = 12
        ws_details.column_dimensions['H'].width = 25
        
        # Format numbers
        for row in ws_details.iter_rows(min_row=2, min_col=4, max_col=5):
            for cell in row:
                cell.number_format = '#,##0.00'
        
        for row in ws_details.iter_rows(min_row=2, min_col=7, max_col=7):
            for cell in row:
                cell.number_format = '0.0'
        
        # Sheet 3: Distribution
        ws_dist = wb.create_sheet("Distribution")
        
        type_counts = {}
        for anom in anomalies:
            anom_type = anom.get('Type', 'Inconnu')
            type_counts[anom_type] = type_counts.get(anom_type, 0) + 1
        
        ws_dist.append(['Type', 'Nombre', 'Pourcentage'])
        
        for cell in ws_dist[1]:
            cell.fill = header_fill
            cell.font = header_font
        
        for anom_type, count in sorted(type_counts.items(), key=lambda x: x[1], reverse=True):
            percentage = (count / len(anomalies)) * 100
            ws_dist.append([anom_type, count, f"{percentage:.1f}%"])
        
        ws_dist.column_dimensions['A'].width = 30
        ws_dist.column_dimensions['B'].width = 15
        ws_dist.column_dimensions['C'].width = 15
        
        # Save to bytes
        buffer = io.BytesIO()
        wb.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()
    
    # ═══════════════════════════════════════════════════════════════
    #                       UTILS
    # ═══════════════════════════════════════════════════════════════
    
    def _avg_confidence(self, anomalies: List[Dict]) -> float:
        """Calculer confiance moyenne"""
        if not anomalies:
            return 0.0
        return sum(a.get('Confiance', 0) for a in anomalies) / len(anomalies)


# ============================================================================
# EXAMPLE USAGE
# ============================================================================

if __name__ == "__main__":
    # Exemple d'utilisation
    generator = ProfessionalAuditReport(company_name="AI_Auditor Test")
    
    # Données test
    test_anomalies = [
        {
            'Type': 'Doublon',
            'Description': 'Doublon détecté: Amazon AWS',
            'Fournisseur': 'Amazon AWS',
            'Montant': 500.00,
            'Impact_Estimation': 49.88,
            'Risque': 'MOYEN',
            'Confiance': 95.0,
            'Recommandation': 'Vérifier les deux factures'
        },
        {
            'Type': 'Écart TPS',
            'Description': 'TPS écart détecté',
            'Fournisseur': 'Bell Canada',
            'Montant': 150.00,
            'Impact_Estimation': 2.50,
            'Risque': 'BAS',
            'Confiance': 85.0,
            'Recommandation': 'Vérifier le numéro TPS'
        }
    ]
    
    # Générer rapports
    for fmt in ['PDF', 'Word', 'Excel']:
        try:
            report_bytes = generator.generate_from_anomalies(
                client_name="Test Client",
                anomalies=test_anomalies,
                format=fmt
            )
            
            filename = f"test_report.{fmt.lower()}"
            with open(filename, 'wb') as f:
                f.write(report_bytes)
            
            print(f"✅ {fmt} rapport généré: {filename}")
        except Exception as e:
            print(f"❌ Erreur {fmt}: {e}")
